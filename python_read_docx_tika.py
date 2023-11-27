# Benchar mark for Tika running on single local machine. Nov.25 2023

import os
import docx
path='/home/ubuntu/chenq/docx_evaluate_score/data/all_docx/input'
inputdir=path
import tika
from tika import parser

def get_text_from_docx(input_docx_path):
    docs = docx.Document(input_docx_path)
    text_list=[]
    for i in range(len(docs.paragraphs)):
        str=docs.paragraphs[i].text
        if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
                text_list.append(str.strip())
    for i in range(len(docs.tables)):
        for j in range(len(docs.tables[i]._cells)):
            str=docs.tables[i]._cells[j].text
            if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
                text_list.append(str.strip())

    #summary = tables[0].rows[6].cells[1].text
    #print(text_list)
    return " ".join(text_list)

ncount=0
print(inputdir)
import time

str_start=time.ctime()
print(str_start)
start = time.time()



for root, dirs, files in os.walk(inputdir):
    for file in files:
        # 获取文件所属目录
        # print(root)
        # 获取文件路径
        # print(os.path.join(root,file))
        if (file.endswith('.docx') and 'output' not in root):
            # print(file+' completed!')
            # image_count=word2pic(file, './tmp', './imgs')
            try:
                #res = get_text_from_docx(inputdir + '/' + file)
                parsed = parser.from_file(inputdir + '/' + file)
            except Exception:
                ncount=ncount
            ncount=ncount+1
        if ncount%100==0:
            
            print(str(ncount)+'th :'+str(int((time.time() - start)))+'s\t'+time.ctime())

print(time.ctime())