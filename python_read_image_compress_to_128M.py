#Compress images to zip file which is no more than 128M each. Updated by Nov.25 2023
import os
import docx
path='/home/ubuntu/chenq/docx_evaluate_score/data/all_docx/input'
#path='/home/ubuntu/chenq/docx_evaluate_score/data/bad_inspect'
#path='/home/ubuntu/chenq/docx_evaluate_score/data/slowest_files'
path='/home/ubuntu/chenq/images/Projection_Shrec14/'
inputdir=path

def get_text_from_docx(input_docx_path):
    docs = docx.Document(input_docx_path)
    text_list=[]
    n_p=len(docs.paragraphs)
    ith=0
    return ""
    if n_p>20000:
        res='can not read file'
        return res
    # for i in range(n_p):
    #     str_=docs.paragraphs[i].text
    #     if str_ not in text_list and str_!=' ' and str_!=' ' and len(str_)>0:
    #             text_list.append(str_.strip())

    for para in docs.paragraphs:
        text_list.append(para.text.strip())

    # ntable=len(docs.tables)
    # for i in range(len(docs.tables)):
    #     #ncell=len(docs.tables[i]._cells)
    #     #print(i)
    #     #print(ncell)
    #     for j in range(len(docs.tables[i]._cells)):
    #         str=docs.tables[i]._cells[j].text
    #         if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
    #             text_list.append(str.strip())
    #             ith=ith+1
    #             if ith % 50 == 0:
    #                 print(ith)



    ith =0
    for table in docs.tables:
        for row in table.rows:
            for cell in row.cells:
                text_list.append(cell.text.strip())
                ith=ith + 1
                if ith%50==0:
                    print(ith)


    return " ".join(text_list)

ncount=0
print(inputdir)
import time
outputdir='/dev/shm/'
str_start=time.ctime()
print(str_start)
start = time.clock()


total_file_size=0 #0M
import zipfile

target='all_images_'

#f = zipfile.ZipFile(target, 'w', zipfile.ZIP_STORED)
#f.write(filename, file_url)
#f.close()

import zipfile


def compress_attaches(files, out_name):
    f = zipfile.ZipFile(out_name, 'w', zipfile.ZIP_STORED)
    for file in files:
        filepath, fullflname = os.path.split(file)
        f.write(file,arcname=filepath.split('/')[-1]+'_'+fullflname)
    f.close()


files_path = [] #['a.txt', 'K.result.xlsx']
#compress_attaches(files_path, target)
n_zip_id=0
for root, dirs, files in os.walk(inputdir):
    for file in files:
        # 
        # print(root)
        # 
        # print(os.path.join(root,file))
        if (file.endswith('.pgm') and 'output' not in root):
            # print(file+' completed!')
            # image_count=word2pic(file, './tmp', './imgs')
            try:
               
                #    a=3
                a=3
                #print(file)
                #res = get_text_from_docx(inputdir + '/' + file)
                #print(len(res))
                #print(file)
                wholepath=root+'/' + file
                files_path.append(wholepath)
                filesize=os.path.getsize(wholepath)/1024/1024


                total_file_size=total_file_size+filesize

                if total_file_size>128:
                    if n_zip_id>0:
                        compress_attaches(files_path, outputdir + '/' + target+str(n_zip_id)+'.zip')
                        print('compressing ' + str(n_zip_id) + ' file.')
                    n_zip_id=n_zip_id+1
                    total_file_size=0
                    files_path.clear()



            except Exception as e:
                ncount=ncount
                print(str(e))
            ncount=ncount+1
            if ncount%1000==0:
                print(ncount)
            
            #print(str(ncount)+'th :'+str(int((time.time() - start)))+'s ' +time.ctime())

if total_file_size > 0:
    compress_attaches(files_path, outputdir + '/' + target + str(n_zip_id) + '.zip')
    n_zip_id = n_zip_id + 1
    print('total file size:' + str(total_file_size) + 'M')
    total_file_size = 0
    files_path.clear()
#Compress images to zip file which is no more than 128M each. Updated by Nov.25 2023

duration=int((time.clock() - start))
print('execution time:\t' + str(duration) + 's')
print(str(ncount)+' files')
