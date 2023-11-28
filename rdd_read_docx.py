#Text processing as depicted in Alogirhtm1,1.1 except tackling 20000 paragraph problem. Nov.25 2023
from pyspark.sql import SparkSession

#from __future__ import print_function

from pyspark import SparkContext
from pyspark import SparkConf

import os

#PYSPARK_PYTHON = "/home/ubuntu/env3.5/bin/python3.5"
#os.environ["PYSPARK_PYTHON"] = PYSPARK_PYTHON



from pyspark.ml.feature import HashingTF,IDF,Tokenizer

import sys
import time
def processing_files_txt(input_docx_path):
	start = time.clock()



	docs = docx.Document(input_docx_path)
	text_list=[]
	n_p =len(docs.paragraphs)
	# if (n_p > 20000):
	# 	res = 'paragraph count >20000'
	# 	print(input_docx_path+' paragraph count >2000  is: '+str(n_p))
	# 	return res
	for i in range(n_p):
		string_=docs.paragraphs[i].text
		#if string_ not in text_list and string_!=' ' and string_!=' ' and len(string_)>0:   #not in too slow
		if string_ != ' ' and string_ != ' ' and len(string_) > 0:
			text_list.append(string_.strip())

	#ith =0
	for table in docs.tables:
	    for row in table.rows:
	        for cell in row.cells:
	            text_list.append(cell.text.strip())   # new method faster than before 190s->24s
	            #ith=ith + 1
	            #if ith%50==0:
	                #print(ith)

	# for i in range(len(docs.tables)):
	# 	for j in range(len(docs.tables[i]._cells)):
	# 		string_=docs.tables[i]._cells[j].text
	# 		#if string_ not in text_list and string_!=' ' and string_!=' ' and len(string_)>0:
	# 		if string_ != ' ' and string_ != ' ' and len(string_) > 0:
	# 			text_list.append(string_.strip())



    #summary = tables[0].rows[6].cells[1].text
    #print(text_list)
	duration=int((time.clock() - start))
	if duration>1:
		print(input_docx_path+' , execution time:\t' + str(duration) + 's')

	return " ".join(text_list)

#spark_conf = SparkConf().setAppName('rdd_read_docx_6_nodes_store_tmp_20cores_16G_tackle_word_cannot_read_problem')
spark_conf = SparkConf()#.setAppName('rdd_read_docx_func_call_testing_115_docs')
#spark_conf = SparkConf().setAppName('rdd_read_docx_no_read_docx')
#spark_conf = SparkConf().setAppName('rdd_read_docx_v4_20cores_7nodes')

    # .setMaster(master)

sc = SparkContext(conf=spark_conf)
#sc.setLogLevel('ERROR')
spark = SparkSession.builder.appName("PythonKMeans").getOrCreate()

total=sc.accumulator(0)

import os
import docx
import shutil
def main_procedure(x,filename):
	#return ""
	filepath, fullflname = os.path.split(filename)
	docfile='/dev/shm/'+fullflname
	docfile = '/tmp/' + fullflname
	res=""
	#print('in func.')
	try:
		with open(docfile, 'wb') as f:
			f.write(x)
			f.close()
			#docs = docx.Document('/tmp/'+filename)
			try:
				res=processing_files_txt(docfile)
				#print('call read func once in shm 200 record time.')
				#print('1')
				#with open(docfile+'_tmp','w') as t:
				#	t.write("")
				#	t.close()
			except Exception as e:
				print('exception: '+str(e))
	except Exception:
		a=3

	#import shutil
	tmp_path=docfile
	os.remove(docfile)
	#print(res)
	print(fullflname+' *optimized table reading in tmp * file length :' + str(len(res)))
	return res
	#for i in os.listdir(tmp_path):
	#	if os.path.isdir(os.path.join(tmp_path, i)):
	#		shutil.rmtree(os.path.join(tmp_path, i))



path='/user/ubuntu/chenq/all_docx/input/*.docx'  #ncount=175420

rdd = sc.binaryFiles(path)

#rdd.cache()
#rdd.count()
#rdd=rdd.repartition(32)

# df=rdd.toDF()
#
# from pyspark.sql.functions import *
# from pyspark.sql.types import *
# from pyspark.sql.types import Row
#
# readUDF=udf(lambda a,z:read(a,z),StringType())
#
# df_new=df.select('_1',readUDF('_2','_1').alias('content'))

doc=rdd.map(lambda x:(x[0],main_procedure(x[1],x[0])))

#doc.foreach(print)
doc.cache()
nread=doc.count()
#df_new.cache()
#nread=df_new.count()
print(nread)


df=doc.toDF()


sc.stop()






