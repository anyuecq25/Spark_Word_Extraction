#
# Algorithm 2. Compress based text processing.  Nov 25,2023
from pyspark.sql import SparkSession

#from __future__ import print_function

from pyspark import SparkContext
from pyspark import SparkConf

import os

#PYSPARK_PYTHON = "/home/ubuntu/env3.5/bin/python3.5"
#os.environ["PYSPARK_PYTHON"] = PYSPARK_PYTHON



from pyspark.ml.feature import HashingTF,IDF,Tokenizer



#spark_conf = SparkConf().setAppName('rdd_read_docx_6_nodes_store_tmp_20cores_16G_tackle_word_cannot_read_problem')
spark_conf = SparkConf()#.setAppName('rdd_read_docx_func_call_testing_115_docs')
#spark_conf = SparkConf().setAppName('rdd_read_docx_no_read_docx')
#spark_conf = SparkConf().setAppName('rdd_read_docx_v4_20cores_7nodes')

    # .setMaster(master)

sc = SparkContext(conf=spark_conf)
#sc.setLogLevel('ERROR')
spark = SparkSession.builder.appName("PythonKMeans").getOrCreate()

total=sc.accumulator(0)

import sys
import time
def processing_files_txt(input_docx_path):
	start = time.clock()
	docs = docx.Document(input_docx_path)
	text_list=[]
	n_p =len(docs.paragraphs)
	if (n_p > 20000):
		res = 'paragraph count >20000'
		print(input_docx_path+' paragraph count >2000  is: '+str(n_p))
		return res
	for i in range(n_p):
		string_=docs.paragraphs[i].text
		if string_ != ' ' and string_ != ' ' and len(string_) > 0:
			text_list.append(string_.strip())
	for table in docs.tables:
	    for row in table.rows:
	        for cell in row.cells:
	            text_list.append(cell.text.strip())   # new method faster than before 190s->24s by Qiang Chen & Yinong Chen et al.
	duration=int((time.clock() - start))
	if duration>1:
		print(input_docx_path+' , execution time:\t' + str(duration) + 's')
	return " ".join(text_list)


import os
import docx
import shutil
def read(filename):
	filepath, fullflname = os.path.split(filename)
	docfile = '/dev/shm/' + filename 	#docfile=filename
	res=""
	try:
		res=processing_files_txt(docfile)
	except Exception as e:
		print('exception: '+str(e))
	os.remove(docfile)
	print(fullflname+' *optimized table reading in zip * file length :' + str(len(res)))
	return res


import io
import zipfile
def main_procedure_zip(content,filename):
	filelist=[]
	itemlist = []
	f = zipfile.ZipFile(io.BytesIO(content))
	for file in f.namelist():  # f.namelist()
		f.extract(file, "/dev/shm/")
		filelist.append(file)
		res=read(file)
		newrow=(file,res)
		itemlist.append(newrow)
	return itemlist


path='/user/ubuntu/chenq/all_docx/input/*.zip'  # ncount=175420 almost 232G by Qiang Chen & Yinong Chen et al. Nov 25, 2023

#export PYSPARK_DRIVER_PYTHON=/home/ubuntu/env3.5/bin/python3.5
#PYSPARK_PYTHON=/home/ubuntu/env3.5/bin/python3.5
#pyspark --master local[12] --executor-memory 32g --driver-memory 10g

rdd = sc.binaryFiles(path)
doc=rdd.flatMap(lambda x:main_procedure_zip(x[1],x[0]))
#doc.cache()
#nread=doc.count()
#print(nread)

df=doc.toDF()
df.cache()
nread=df.count()
print(nread)
import time
ticks = time.time()
saved_parquet_name='./mini_zip_docx/mini_zip_docx_'+str(int(ticks))+'.parquet'
print('saved_parquet_name :'+saved_parquet_name)
df.write.parquet(saved_parquet_name)

sc.stop()




