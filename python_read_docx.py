import os
import docx
path='/home/ubuntu/chenq/docx_evaluate_score/data/all_docx/input'
#path='/home/ubuntu/chenq/docx_evaluate_score/data/bad_inspect'
#path='/home/ubuntu/chenq/docx_evaluate_score/data/slowest_files'
inputdir=path

def get_text_from_docx(input_docx_path):
    docs = docx.Document(input_docx_path)
    text_list=[]
    n_p=len(docs.paragraphs)
    ith=0
    #return ""
    if n_p>20000:
        res='The No. of paragraphs is greater than 20000'
        return res
    # for i in range(n_p):
    #     str_=docs.paragraphs[i].text
    #     if str_ not in text_list and str_!=' ' and str_!=' ' and len(str_)>0:
    #             text_list.append(str_.strip())

    for para in docs.paragraphs:
        text_list.append(para.text.strip())

    # ntable=len(docs.tables)
    # for i in range(len(docs.tables)):
    #     #ncell=len(docs.tables[i]._cells)
    #     #print(i)
    #     #print(ncell)
    #     for j in range(len(docs.tables[i]._cells)):
    #         str=docs.tables[i]._cells[j].text
    #         if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
    #             text_list.append(str.strip())
    #             ith=ith+1
    #             if ith % 50 == 0:
    #                 print(ith)



    ith =0
    for table in docs.tables:
        for row in table.rows:
            for cell in row.cells:
                text_list.append(cell.text.strip())
                ith=ith + 1
                #if ith%50==0:
                    #print(ith)


    return " ".join(text_list)

ncount=0
print(inputdir)
import time

str_start=time.ctime()
print(str_start)
start = time.clock()



for root, dirs, files in os.walk(inputdir):
    for file in files:
        # 获取文件所属目录
        # print(root)
        # 获取文件路径
        # print(os.path.join(root,file))
        if (file.endswith('.docx') and 'output' not in root):
            # print(file+' completed!')
            # image_count=word2pic(file, './tmp', './imgs')
            try:
                #if file=='222016310011059许文欣邮件合并.docx': may cost 1479s
                #    a=3
                a=3
                #print(file)
                res = get_text_from_docx(inputdir + '/' + file)
                if len(res)==42:
                    print(res+'. '+file+' .The '+str(ncount))
                else:
                    print('File length :'+str(len(res))+' '+file+' .The '+str(ncount))
            except Exception as e:
                print('Exception in' +str(e)+' .The '+str(ncount))
                ncount=ncount
            ncount=ncount+1
        #if ncount%10==0:
            
            #print(str(ncount)+'th :'+str(int((time.time() - start)))+'s ' +time.ctime())
#Benchmark for Python-docx. Running on one thread. Serial algorithm. Updated by Nov.25 2023
duration=int((time.clock() - start))
print('execution time:\t' + str(duration) + 's')
print('Total file is :'+str(ncount))

str_start=time.ctime()
print(str_start)