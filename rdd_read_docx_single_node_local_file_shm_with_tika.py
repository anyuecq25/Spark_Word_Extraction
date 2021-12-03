from pyspark.sql import SparkSession

#from __future__ import print_function

from pyspark import SparkContext
from pyspark import SparkConf

import os

#PYSPARK_PYTHON = "/home/ubuntu/env3.5/bin/python3.5"
#os.environ["PYSPARK_PYTHON"] = PYSPARK_PYTHON



from pyspark.ml.feature import HashingTF,IDF,Tokenizer

import sys
import tika
from tika import parser

def get_text_from_docx(input_docx_path):
    docs = docx.Document(input_docx_path)
    text_list=[]
    for i in range(len(docs.paragraphs)):
        str=docs.paragraphs[i].text
        if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
                text_list.append(str.strip())
    for i in range(len(docs.tables)):
        for j in range(len(docs.tables[i]._cells)):
            str=docs.tables[i]._cells[j].text
            if str not in text_list and str!=' ' and str!=' ' and len(str)>0:
                text_list.append(str.strip())

    #summary = tables[0].rows[6].cells[1].text
    #print(text_list)
    return " ".join(text_list)


spark_conf = SparkConf() \
        .setAppName('rdd_read_docx_only_local_file_10349_files__on_single_node_read_shm_24_cores_tika')
    # .setMaster(master)

sc = SparkContext(conf=spark_conf)
#sc.setLogLevel('ERROR')
spark = SparkSession.builder.appName("rdd_read_docx_only_local_file_10349_files__on_single_node_read_shm_").getOrCreate()

total=sc.accumulator(0)

import os
import docx
import shutil
def read(x,filename):
    filepath, fullflname = os.path.split(filename)
    docfile='/dev/shm/'+fullflname
    res=""
    with open(docfile, 'wb') as f:
        f.write(x)
        f.close()
        try:
            res = parser.from_file(docfile)

        except Exception:
            a=2
    tmp_path=docfile
    os.remove(docfile)
    return res
	#for i in os.listdir(tmp_path):
	#	if os.path.isdir(os.path.join(tmp_path, i)):
	#		shutil.rmtree(os.path.join(tmp_path, i))


path='file:///home/ubuntu/chenq/docx_evaluate_score/data/all_docx/input/*.docx'
path='file:///dev/shm/test_docx/input/*.docx'
#path='file:///home/ubuntu/chenq/test_docx/input/*.docx'
#path='file:////dev/shm/input/*.docx'
#path='file:////dev/shm/input2/input/*.docx'
#path='/user/ubuntu/docx/input/*.docx'
#path='har:////user/ubuntu/ainput/docx.har/input/*.docx'
rdd = sc.binaryFiles(path)
#rdd.cache()
#rdd.count()
#rdd=rdd.repartition(32)
doc=rdd.map(lambda x:(x[0],read(x[1],x[0])))

#doc.foreach(print)
doc.cache()
nread=doc.count()

df=doc.toDF()
df.write.parquet('file:///user/ubuntu/all_docx_10349_single_node_on_desktop5_localfile_shm_24cores_tika.parquet')


print(nread)
sc.stop()
'''
while(nread>0):
	rdd1 = sc.binaryFiles(path)
	doc1 = rdd1.map(lambda x: (x[0], read(x[1], x[0])))
	nread = doc1.count()
	if(nread>0):
		doc=doc.union(doc1)
		print('This time read: '+str(nread))

print('total count is:'+str(doc.count()))
#print(doc.first()[0])
#print(doc.first()[1])
'''





